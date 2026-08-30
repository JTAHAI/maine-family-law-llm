from __future__ import annotations

import argparse
import hashlib
import io
import ipaddress
import json
import os
import re
import socket
import subprocess
import sys
import tempfile
import threading
import time
import urllib.error
import urllib.parse
import urllib.request
import uuid
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
if str(ROOT / "src") not in sys.path:
    sys.path.insert(0, str(ROOT / "src"))

from docx import Document
from fpdf import FPDF
from PIL import Image, ImageDraw, ImageFont
from pypdf import PdfReader

from maine_family_law_llm.installed_runtime import (
    DEFAULT_PACKAGE_NAME,
    InstalledRuntimeResolution,
    resolve_installed_runtime_executable,
)
from maine_family_law_llm.local_only_boundary import LocalOnlyNetworkBlocked, local_only_network_boundary

QA_HEADERS = {
    "X-User-Role": "reviewer",
    # Use the production desktop tenant for this isolated fictional profile so
    # API-created audit ownership remains compatible with later real UI actions.
    "X-Tenant-Id": "local-desktop",
    "X-MFLL-Client-Session": uuid.uuid4().hex + uuid.uuid4().hex,
}
REQUEST_EVENTS: list[dict[str, Any]] = []


def utc_now() -> str:
    return datetime.now(UTC).replace(microsecond=0).isoformat().replace("+00:00", "Z")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def free_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        sock.bind(("127.0.0.1", 0))
        return int(sock.getsockname()[1])


def wait_json(url: str, *, timeout_s: int = 120) -> dict[str, Any]:
    deadline = time.time() + timeout_s
    last_error: str = "not_started"
    while time.time() < deadline:
        try:
            with urllib.request.urlopen(url, timeout=5) as response:
                return json.loads(response.read().decode("utf-8"))
        except Exception as exc:  # noqa: BLE001
            last_error = f"{exc.__class__.__name__}: {exc}"
            time.sleep(1.5)
    raise TimeoutError(f"Timed out waiting for {url}: {last_error}")


def request_json(method: str, url: str, payload: dict[str, Any] | None = None) -> dict[str, Any]:
    body = None
    headers = {**QA_HEADERS, "Accept": "application/json"}
    if payload is not None:
        body = json.dumps(payload).encode("utf-8")
        headers["Content-Type"] = "application/json"
    request = urllib.request.Request(url, data=body, method=method.upper(), headers=headers)
    started = time.monotonic()
    event = {"method": method.upper(), "path": urllib.parse.urlsplit(url).path}
    try:
        with urllib.request.urlopen(request, timeout=660) as response:
            event["http_status"] = response.status
            event["service_instance"] = response.headers.get("X-MFL-Service-Instance", "")
            event["request_id"] = response.headers.get("X-Request-ID", "")
            raw = response.read().decode("utf-8")
            return json.loads(raw) if raw else {}
    except Exception as exc:
        event["error_class"] = type(exc).__name__
        event["http_status"] = getattr(exc, "code", None)
        raise
    finally:
        event["duration_seconds"] = round(time.monotonic() - started, 3)
        REQUEST_EVENTS.append(event)
        print(json.dumps({"request": event["path"], "status": event.get("http_status"),
                          "duration_seconds": event["duration_seconds"]}), flush=True)


def download_bytes(url: str) -> bytes:
    with urllib.request.urlopen(urllib.request.Request(url, headers=QA_HEADERS), timeout=120) as response:
        return response.read()


def verify_runtime_instance(base_url: str, instance: str) -> bool:
    if not re.fullmatch(r"[0-9a-f]{64}", instance):
        return False
    try:
        call = urllib.request.Request(base_url + "/api/health", headers=QA_HEADERS)
        with urllib.request.urlopen(call, timeout=5) as response:
            return response.status == 200 and response.headers.get("X-MFL-Service-Instance") == instance
    except (OSError, urllib.error.URLError):
        return False


def verified_artifact_bytes(base_url: str, artifact: dict[str, Any]) -> bytes:
    route = str(artifact.get("download_url") or "")
    if not re.fullmatch(r"/api/document-intelligence/artifacts/[a-f0-9]{64}", route):
        raise ValueError("artifact_capability_route_required")
    size = artifact.get("size_bytes")
    if type(size) is not int or not 0 < size <= 64 * 1024 * 1024:
        raise ValueError("artifact_size_invalid")
    call = urllib.request.Request(base_url + route, headers=QA_HEADERS)
    with urllib.request.urlopen(call, timeout=120) as response:
        if response.headers.get("X-MFL-Hash-Verified") != "true":
            raise ValueError("artifact_not_verified_by_runtime")
        data = response.read(size + 1)
    if len(data) != size or hashlib.sha256(data).hexdigest() != artifact.get("sha256"):
        raise ValueError("artifact_content_binding_failed")
    return data


def verified_pdf_page(base_url: str, metadata: dict[str, Any], page: int, *, instance: str) -> tuple[bytes, dict[str, Any]]:
    token = str(metadata.get("token") or "")
    if not re.fullmatch(r"[a-f0-9]{64}", token) or type(page) is not int or page < 1:
        raise ValueError("pdf_preview_capability_required")
    route = f"/api/records/preview/{token}?page={page}"
    started = time.monotonic()
    event = {"method": "GET", "path": "/api/records/preview/{token}", "page": page}
    try:
        if not verify_runtime_instance(base_url, instance):
            raise ValueError("pdf_preview_runtime_instance_unverified")
        call = urllib.request.Request(base_url + route, headers=QA_HEADERS)
        with urllib.request.urlopen(call, timeout=45) as response:
            headers = response.headers
            event["http_status"] = response.status
            event["service_instance"] = headers.get("X-MFL-Service-Instance", "")
            if (response.status != 200 or headers.get_content_type() != "image/png"
                    or headers.get("X-MFL-Hash-Verified") != "true"
                    or headers.get("X-MFL-Source-Hash") != metadata.get("source_hash")
                    or headers.get("X-MFL-Page") != str(page)
                    or headers.get("X-MFL-Review-Required") != "true"
                    or not re.fullmatch(r"[a-f0-9]{64}", headers.get("X-MFL-Audit-Receipt", ""))
                    or (event["service_instance"] and event["service_instance"] != instance)):
                raise ValueError("pdf_preview_binding_failed")
            data = response.read(8 * 1024 * 1024 + 1)
            if len(data) > 8 * 1024 * 1024 or hashlib.sha256(data).hexdigest() != headers.get("X-MFL-Preview-Hash"):
                raise ValueError("pdf_preview_hash_failed")
            page_count = int(headers.get("X-MFL-Page-Count", "0"))
            with Image.open(io.BytesIO(data)) as image:
                image.load()
                if image.format != "PNG" or not page <= page_count <= 100000 or not 1 <= min(image.size) <= max(image.size) <= 1600:
                    raise ValueError("pdf_preview_raster_invalid")
                dimensions = list(image.size)
            # The secret instance header is intentionally exposed only by the
            # health route, not by record responses. Bind this action on both
            # sides using the same loopback origin and owned runtime instance.
            if not verify_runtime_instance(base_url, instance):
                raise ValueError("pdf_preview_runtime_instance_changed")
            event["runtime_instance_verified_before_and_after"] = True
            return data, {"status": "pass", "page": page, "page_count": page_count,
                          "dimensions": dimensions, "sha256": hashlib.sha256(data).hexdigest(),
                          "source_hash": metadata["source_hash"], "review_required": True,
                          "audit_receipt": headers["X-MFL-Audit-Receipt"]}
    finally:
        event["duration_seconds"] = round(time.monotonic() - started, 3)
        REQUEST_EVENTS.append(event)


def ocr_completed(result: dict[str, Any], pdf_text: str, sidecar: bytes) -> bool:
    """Blocked, empty or unreviewed derivatives never qualify as successful OCR."""
    phrase = "scan for ocr"
    return (
        result.get("status") == "pass"
        and not result.get("blockers")
        and result.get("review_required") is True
        and result.get("original_modified") is False
        and phrase in " ".join(pdf_text.lower().split())
        and phrase in " ".join(sidecar.decode("utf-8", errors="replace").lower().split())
    )


def create_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Motion for Temporary Relief", level=1)
    document.add_paragraph("The child changed schools on January 3, 2026.")
    document.add_paragraph("Parent Jane Example can be reached at jane.example@example.com.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Date"
    table.cell(0, 1).text = "Event"
    table.cell(1, 0).text = "2026-01-03"
    table.cell(1, 1).text = "School change"
    document.save(path)


def create_pdf_with_text(path: Path, lines: list[str], *, second_page: bool = False) -> None:
    pdf = FPDF(unit="pt", format="letter")
    pdf.add_page()
    pdf.set_font("Helvetica", size=14)
    y = 48
    for line in lines:
        pdf.set_xy(48, y)
        pdf.multi_cell(516, 18, line)
        y += 24
    if second_page:
        pdf.add_page()
        pdf.set_xy(48, 48)
        pdf.multi_cell(516, 22, "FICTIONAL PAGE TWO - Review required. No real records.")
        pdf.set_fill_color(20, 120, 130)
        pdf.rect(48, 140, 300, 120, style="F")
    pdf.output(str(path))


def create_image_only_pdf(path: Path, *, label: str) -> Path:
    image_path = path.with_suffix(".png")
    image = Image.new("RGB", (1240, 1754), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default(size=36)
    draw.text((80, 120), f"FICTIONAL {label}", fill="black", font=font)
    draw.text((80, 180), "Review required. No real records.", fill="black", font=font)
    image.save(image_path)
    pdf = FPDF(unit="pt", format=(612, 792))
    pdf.add_page()
    pdf.image(str(image_path), x=36, y=36, w=540)
    pdf.output(str(path))
    return image_path


def stage_record(case_root: Path, *, evidence_id: str, filename: str, data: bytes, **row: Any) -> dict[str, Any]:
    rel_path = Path("02_PRIVATE_FORENSIC_MASTER") / "files" / filename
    target = case_root / rel_path
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_bytes(data)
    source_hash = hashlib.sha256(data).hexdigest()
    return {
        "evidence_id": evidence_id,
        "title": row.get("title") or evidence_id,
        "source_type": row.get("source_type") or target.suffix.lstrip("."),
        "source_locator": rel_path.as_posix(),
        "private_copy_relpath": rel_path.as_posix(),
        "source_hash": source_hash,
        "page_number": int(row.get("page_number") or 1),
        "page_count": int(row.get("page_count") or 1),
        "parser_status": row.get("parser_status") or "parsed",
        "text_status": row.get("text_status") or "available",
        "ocr_status": row.get("ocr_status") or "not_run",
        "text_excerpt": str(row.get("text_excerpt") or ""),
        "text_content": str(row.get("text_content") or ""),
        "issue_lanes": list(row.get("issue_lanes") or []),
        "procedural_postures": list(row.get("procedural_postures") or []),
    }


def build_case_fixture(case_root: Path) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    manifest: list[dict[str, Any]] = []

    pii_text = (
        "FICTIONAL SOFTWARE TEST. Motion for temporary relief.\n"
        "Parent Fictional Example lives at 10 Example Lane, Fictional Town.\n"
        "Email jane.example@example.com and phone (207) 555-0109.\n"
        "Fictional DOB 01/02/2010 and invalid SSN 000-00-0000 are test data."
    )
    pii_txt = case_root / "pii.txt"
    pii_txt.write_text(pii_text, encoding="utf-8")
    manifest.append(
        {
            "evidence_id": "REC-PII-TXT",
            "source_path": str(pii_txt.resolve()),
            "private_copy_relpath": (Path("02_PRIVATE_FORENSIC_MASTER/files") / pii_txt.name).as_posix(),
            "source_hash": hashlib.sha256(pii_txt.read_bytes()).hexdigest(),
            "subject": "PII note",
        }
    )
    records.append(
        stage_record(
            case_root,
            evidence_id="REC-PII-TXT",
            filename=pii_txt.name,
            data=pii_txt.read_bytes(),
            title="PII note",
            text_excerpt=pii_text,
            text_content=pii_text,
            issue_lanes=["privacy", "contact"],
        )
    )

    docx_path = case_root / "motion.docx"
    create_docx(docx_path)
    manifest.append(
        {
            "evidence_id": "REC-DOCX",
            "source_path": str(docx_path.resolve()),
            "private_copy_relpath": (Path("02_PRIVATE_FORENSIC_MASTER/files") / docx_path.name).as_posix(),
            "source_hash": hashlib.sha256(docx_path.read_bytes()).hexdigest(),
            "subject": "Motion for Temporary Relief",
        }
    )
    records.append(
        stage_record(
            case_root,
            evidence_id="REC-DOCX",
            filename=docx_path.name,
            data=docx_path.read_bytes(),
            title="Motion for Temporary Relief",
            text_excerpt="The child changed schools on January 3, 2026.",
            text_content="The child changed schools on January 3, 2026. Parent Jane Example can be reached at jane.example@example.com.",
            issue_lanes=["custody", "school"],
        )
    )

    pdf_path = case_root / "brief.pdf"
    create_pdf_with_text(
        pdf_path,
        [
            "This ordinary PDF has a heading, a body paragraph, and a small table.",
            "The child changed schools on January 3, 2026.",
            "Contact: jane.example@example.com",
        ],
        second_page=True,
    )
    manifest.append(
        {
            "evidence_id": "REC-PDF",
            "source_path": str(pdf_path.resolve()),
            "private_copy_relpath": (Path("02_PRIVATE_FORENSIC_MASTER/files") / pdf_path.name).as_posix(),
            "source_hash": hashlib.sha256(pdf_path.read_bytes()).hexdigest(),
            "subject": "Ordinary PDF",
        }
    )
    records.append(
        stage_record(
            case_root,
            evidence_id="REC-PDF",
            filename=pdf_path.name,
            data=pdf_path.read_bytes(),
            title="Ordinary PDF",
            page_count=2,
            text_excerpt="The child changed schools on January 3, 2026.",
            text_content="This ordinary PDF has a heading, a body paragraph, and a small table. The child changed schools on January 3, 2026.",
            issue_lanes=["school"],
        )
    )

    ocr_pdf = case_root / "scan.pdf"
    create_image_only_pdf(ocr_pdf, label="Scan for OCR")
    manifest.append(
        {
            "evidence_id": "REC-OCR",
            "source_path": str(ocr_pdf.resolve()),
            "private_copy_relpath": (Path("02_PRIVATE_FORENSIC_MASTER/files") / ocr_pdf.name).as_posix(),
            "source_hash": hashlib.sha256(ocr_pdf.read_bytes()).hexdigest(),
            "subject": "Image-only scan",
        }
    )
    records.append(
        stage_record(
            case_root,
            evidence_id="REC-OCR",
            filename=ocr_pdf.name,
            data=ocr_pdf.read_bytes(),
            title="Image-only scan",
            text_excerpt="Scan for OCR",
            text_content="Scan for OCR",
            source_type="pdf",
            issue_lanes=["ocr"],
        )
    )

    image_path = ocr_pdf.with_suffix(".png")
    image_bytes = image_path.read_bytes()
    manifest.append({
        "evidence_id": "REC-IMAGE", "source_path": str(image_path.resolve()),
        "private_copy_relpath": "02_PRIVATE_FORENSIC_MASTER/files/scan.png",
        "source_hash": hashlib.sha256(image_bytes).hexdigest(),
        "subject": "Fictional image scan for OCR",
    })
    records.append(stage_record(case_root, evidence_id="REC-IMAGE", filename="scan.png",
                                data=image_bytes, title="Fictional image scan for OCR"))

    dup_a = case_root / "duplicate-a.docx"
    create_docx(dup_a)
    dup_b = case_root / "duplicate-b.docx"
    # Copy the same bytes so the qualification fixture is an exact duplicate;
    # independently saving two DOCX files can embed different ZIP timestamps.
    dup_b.write_bytes(dup_a.read_bytes())
    manifest.append(
        {
            "evidence_id": "REC-DUP-A",
            "source_path": str(dup_a.resolve()),
            "private_copy_relpath": (Path("02_PRIVATE_FORENSIC_MASTER/files") / dup_a.name).as_posix(),
            "source_hash": hashlib.sha256(dup_a.read_bytes()).hexdigest(),
            "subject": "Duplicate A",
        }
    )
    manifest.append(
        {
            "evidence_id": "REC-DUP-B",
            "source_path": str(dup_b.resolve()),
            "private_copy_relpath": (Path("02_PRIVATE_FORENSIC_MASTER/files") / dup_b.name).as_posix(),
            "source_hash": hashlib.sha256(dup_b.read_bytes()).hexdigest(),
            "subject": "Duplicate B",
        }
    )
    dup_text = "The child changed schools on January 3, 2026."
    records.append(
        stage_record(
            case_root,
            evidence_id="REC-DUP-A",
            filename=dup_a.name,
            data=dup_a.read_bytes(),
            title="Duplicate A",
            text_excerpt=dup_text,
            text_content=dup_text,
            issue_lanes=["duplication"],
        )
    )
    records.append(
        stage_record(
            case_root,
            evidence_id="REC-DUP-B",
            filename=dup_b.name,
            data=dup_b.read_bytes(),
            title="Duplicate B",
            text_excerpt=dup_text,
            text_content=dup_text,
            issue_lanes=["duplication"],
        )
    )

    changed = case_root / "changed-copy.docx"
    changed_doc = Document(dup_a)
    changed_doc.add_paragraph("FICTIONAL changed copy: a lighthouse attachment is still missing.")
    changed_doc.save(changed)
    changed_bytes = changed.read_bytes()
    manifest.append({
        "evidence_id": "REC-CHANGED", "source_path": str(changed.resolve()),
        "private_copy_relpath": "02_PRIVATE_FORENSIC_MASTER/files/changed-copy.docx",
        "source_hash": hashlib.sha256(changed_bytes).hexdigest(),
        "subject": "Fictional changed copy",
    })
    records.append(stage_record(
        case_root, evidence_id="REC-CHANGED", filename=changed.name, data=changed_bytes,
        title="Fictional changed copy", text_excerpt=dup_text + " Lighthouse attachment missing.",
        text_content=dup_text + " Lighthouse attachment missing.",
    ))

    index_root = case_root / "04_INDEXES"
    index_root.mkdir(parents=True, exist_ok=True)
    (index_root / "private_search_index.json").write_text(json.dumps(records, indent=2, sort_keys=True), encoding="utf-8")
    manifest_root = case_root / "08_SOURCE_MANIFESTS_HASHES"
    manifest_root.mkdir(parents=True, exist_ok=True)
    (manifest_root / "source_manifest.json").write_text(json.dumps(manifest, indent=2, sort_keys=True), encoding="utf-8")
    return records


def _normalized_text(value: str) -> str:
    return " ".join(value.lower().split())


def _duplicate_report(rows: list[dict[str, Any]], record_id: str) -> dict[str, Any]:
    by_id = {str(row.get("evidence_id") or ""): dict(row) for row in rows}
    row = dict(by_id.get(record_id) or {})
    target_hash = str(row.get("source_hash") or "")
    exact = [dict(item) for item in rows if str(item.get("source_hash") or "") == target_hash and target_hash]
    base_text = _normalized_text(str(row.get("text_excerpt") or row.get("text_content") or ""))
    near_candidates = []
    for item in rows:
        other_text = _normalized_text(str(item.get("text_excerpt") or item.get("text_content") or ""))
        if not other_text:
            continue
        similarity = 1.0 if other_text == base_text else float(__import__("difflib").SequenceMatcher(None, base_text, other_text).ratio())
        if similarity >= 0.7 and str(item.get("evidence_id") or "") != record_id:
            near_candidates.append(
                {
                    "evidence_id": item.get("evidence_id"),
                    "similarity": round(similarity, 6),
                    "source_hash": item.get("source_hash"),
                }
            )
    return {
        "schema_version": "record_duplicate_report_v1",
        "record_id": record_id,
        "duplicate_group_id": target_hash or record_id,
        "exact_duplicate": len(exact) > 1,
        "exact_duplicates": exact,
        "near_duplicate_candidates": near_candidates[:100],
    }


def _record_compare(left: dict[str, Any], right: dict[str, Any]) -> dict[str, Any]:
    left_text = _normalized_text(str(left.get("text_excerpt") or left.get("text_content") or ""))
    right_text = _normalized_text(str(right.get("text_excerpt") or right.get("text_content") or ""))
    similarity = __import__("difflib").SequenceMatcher(None, left_text, right_text).ratio() if left_text or right_text else 0.0
    field_differences: dict[str, dict[str, Any]] = {}
    for field_name in ("source_hash", "page_count", "parser_status", "ocr_status", "text_status", "source_type", "canonical_document_key"):
        if str(left.get(field_name) or "") != str(right.get(field_name) or ""):
            field_differences[field_name] = {"left": left.get(field_name), "right": right.get(field_name)}
    return {
        "schema_version": "record_compare_response_v1",
        "left_record_id": str(left.get("evidence_id") or ""),
        "right_record_id": str(right.get("evidence_id") or ""),
        "same": str(left.get("source_hash") or "") == str(right.get("source_hash") or ""),
        "exact_duplicate": str(left.get("source_hash") or "") == str(right.get("source_hash") or ""),
        "similarity": round(similarity, 6),
        "page_count_delta": int(right.get("page_count") or 0) - int(left.get("page_count") or 0),
        "text_additions": sorted(set(right_text.split()) - set(left_text.split()))[:100],
        "text_removals": sorted(set(left_text.split()) - set(right_text.split()))[:100],
        "field_differences": field_differences,
        "left": {
            "record_id": str(left.get("evidence_id") or ""),
            "source_hash": left.get("source_hash"),
            "page_count": left.get("page_count"),
            "parser_status": left.get("parser_status"),
            "ocr_status": left.get("ocr_status"),
        },
        "right": {
            "record_id": str(right.get("evidence_id") or ""),
            "source_hash": right.get("source_hash"),
            "page_count": right.get("page_count"),
            "parser_status": right.get("parser_status"),
            "ocr_status": right.get("ocr_status"),
        },
    }


def start_runtime(
    executable: Path,
    port: int,
    *,
    localappdata: Path,
    authority_data_root: Path | None = None,
    transfer_root: Path | None = None,
) -> subprocess.Popen[str]:
    """Start an isolated frozen runtime.

    Qualification defaults to an empty authority root so its Local-only proof
    cannot accidentally borrow a developer's external authority store.  A
    focused source-bound workflow may opt in to a caller-supplied, separately
    audited external root; it remains outside both the temporary profile and
    the MSIX.
    """

    environment = os.environ.copy()
    environment.pop("MAINE_MATTER_STORE_KEY", None)
    for key in list(environment):
        if key.startswith("MFL_FAST_INTERCHANGE_") or key == "MAINE_FAST_INTERCHANGE_WORKER_TOKEN":
            environment.pop(key)
    instance = uuid.uuid4().hex + uuid.uuid4().hex
    environment.update(
        {
            "LOCALAPPDATA": str(localappdata),
            "MAINE_FAMILY_LAW_DATA_ROOT": str(localappdata / "runtime"),
            "MFL_AUTHORITY_DATA_ROOT": str(authority_data_root or (localappdata / "empty-authority")),
            "MFL_RUNTIME_STATE_ROOT": str(localappdata / "state"),
            "MFL_IDEMPOTENCY_STATE_ROOT": str(localappdata / "idempotency"),
            "MFL_VAULT_KEY_ROOT": str(localappdata / "vault"),
            # Cross-device transfer is intentionally unavailable unless the
            # user selects an external, user-carried destination.  The
            # qualification harness supplies a disposable sibling directory
            # only for the focused recovery proof; ordinary qualification
            # keeps this unset and therefore fail-closed.
            "MFL_TRANSFER_ROOT": str(transfer_root) if transfer_root is not None else "",
            "MFL_LOCAL_API_INSTANCE_ID": instance,
            "PYTHONDONTWRITEBYTECODE": "1",
            "PYTHONPYCACHEPREFIX": str(localappdata / "bytecode"),
            "MFL_RUNTIME_MODE": "store",
            "HTTP_PROXY": "http://127.0.0.1:9",
            "HTTPS_PROXY": "http://127.0.0.1:9",
            "ALL_PROXY": "http://127.0.0.1:9",
            "NO_PROXY": "127.0.0.1,localhost,::1",
            "HF_HUB_OFFLINE": "1",
            "TRANSFORMERS_OFFLINE": "1",
            "TLD_EXTRACT_NO_FETCH": "1",
        }
    )
    process = subprocess.Popen(
        [str(executable), "--serve-local-api", "--port", str(port)],
        cwd=str(executable.parent),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        text=True,
        env=environment,
        creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
    )
    process.qa_service_instance = instance
    return process


class RuntimeNetworkMonitor:
    """Best-effort OS observation of runtime and worker TCP connections."""

    def __init__(self, root_pid: int, *, interval_seconds: float = 0.025) -> None:
        self.root_pid = int(root_pid)
        self.interval_seconds = float(interval_seconds)
        self.stop_event = threading.Event()
        self.thread = threading.Thread(target=self._run, name="mfl-runtime-network-monitor", daemon=True)
        self.observations: list[dict[str, Any]] = []
        self.errors: list[str] = []
        self.sample_count = 0

    @staticmethod
    def _loopback(host: str) -> bool:
        try:
            return ipaddress.ip_address(host.split("%", 1)[0]).is_loopback
        except ValueError:
            return host.lower() in {"localhost"}

    def start(self) -> None:
        self.thread.start()

    def stop(self) -> dict[str, Any]:
        self.stop_event.set()
        self.thread.join(timeout=10)
        unique = {
            (row["pid"], row["remote_host"], row["remote_port"], row["status"]): row
            for row in self.observations
        }
        rows = sorted(unique.values(), key=lambda row: (row["pid"], row["remote_host"], row["remote_port"], row["status"]))
        external = [row for row in rows if not row["loopback"]]
        return {
            "method": "psutil_process_tree_tcp_polling",
            "interval_ms": round(self.interval_seconds * 1000, 3),
            "sample_count": self.sample_count,
            "observed_remote_connections": rows,
            "external_connections": external,
            "external_connection_count": len(external),
            "errors": self.errors,
            "proxy_fail_closed": True,
            "offline_library_flags": ["HF_HUB_OFFLINE", "TRANSFORMERS_OFFLINE", "TLD_EXTRACT_NO_FETCH"],
        }

    def _run(self) -> None:
        try:
            import psutil
        except Exception as exc:  # noqa: BLE001
            self.errors.append(f"psutil_unavailable:{exc.__class__.__name__}")
            return
        while not self.stop_event.is_set():
            self.sample_count += 1
            try:
                root = psutil.Process(self.root_pid)
                processes = [root, *root.children(recursive=True)]
                for process in processes:
                    try:
                        connections = process.net_connections(kind="tcp")
                    except psutil.AccessDenied:
                        marker = "process_connections_access_denied"
                        if marker not in self.errors:
                            self.errors.append(marker)
                        continue
                    except psutil.NoSuchProcess:
                        continue
                    for connection in connections:
                        if not connection.raddr:
                            continue
                        host = str(connection.raddr.ip)
                        self.observations.append(
                            {
                                "pid": int(process.pid),
                                "remote_host": host,
                                "remote_port": int(connection.raddr.port),
                                "status": str(connection.status),
                                "loopback": self._loopback(host),
                            }
                        )
            except psutil.NoSuchProcess:
                break
            except Exception as exc:  # noqa: BLE001
                marker = exc.__class__.__name__
                if marker not in self.errors:
                    self.errors.append(marker)
            self.stop_event.wait(self.interval_seconds)


def _runtime_resolution(explicit_executable: str) -> InstalledRuntimeResolution:
    if not explicit_executable:
        return resolve_installed_runtime_executable()
    executable = Path(explicit_executable).expanduser().resolve(strict=True)
    if not executable.is_file():
        raise SystemExit(f"Runtime executable is not a regular file: {executable}")
    return InstalledRuntimeResolution(
        package_name=DEFAULT_PACKAGE_NAME,
        package_full_name="",
        install_location=str(executable.parent),
        version="",
        executable_path=str(executable),
        source="explicit_bundled_runtime",
        available=True,
    )


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--runtime-executable",
        default="",
        help="Qualify this exact frozen runtime instead of resolving an installed AppX package.",
    )
    parser.add_argument("--evidence-root", type=Path)
    parser.add_argument("--hold-seconds", type=int, default=0,
                        help="Keep this fictional runtime available for browser checks; stop with stop-probe in the evidence directory.")
    args = parser.parse_args(argv)
    if not 0 <= args.hold_seconds <= 3600:
        parser.error("hold-seconds must be between 0 and 3600")
    evidence_root = Path(
        args.evidence_root or os.environ.get("MFL_GA_EVIDENCE_ROOT")
        or ROOT / "dist" / "store" / "evidence"
    ).expanduser().resolve()
    evidence_root.mkdir(parents=True, exist_ok=True)
    if (evidence_root / "installed-offline-qualification.json").exists():
        raise SystemExit("Use a new evidence directory; existing qualification evidence is immutable.")
    REQUEST_EVENTS.clear()
    resolution = _runtime_resolution(args.runtime_executable)
    if not resolution.executable_path:
        raise SystemExit("Unable to resolve an installed or bundled runtime executable.")

    port = free_port()
    base_url = f"http://127.0.0.1:{port}"
    summary: dict[str, Any] = {
        "schema_version": "installed_offline_qualification_v2",
        "generated_at": utc_now(),
        "runtime_resolution": resolution.as_dict(),
        "runtime_sha256": sha256_file(Path(resolution.executable_path)),
        "execution_level": "frozen_runtime_canonical_http",
        "installed_msix": resolution.source == "appx_package",
        "fictional_only": True,
        "encryption_path": "isolated_production_managed_vault_no_environment_matter_key",
        "browser_interaction_tested": False,
        "base_url": base_url,
        "offline_boundary": {},
        "feature_results": {},
        "qualification_checks": {},
        "inventory_result": {},
        "qualification_status": "blocked",
        "feature_check_status": "blocked",
        "blockers": [],
    }

    original_socket = socket.socket
    original_create_connection = socket.create_connection
    original_getaddrinfo = socket.getaddrinfo

    with local_only_network_boundary():
        try:
            urllib.request.urlopen("https://example.com", timeout=5)
            summary["offline_boundary"]["external_request"] = "unexpected_success"
            summary["blockers"].append("offline_boundary_did_not_block_external_request")
        except LocalOnlyNetworkBlocked as exc:
            summary["offline_boundary"]["external_request"] = "blocked"
            summary["offline_boundary"]["detail"] = str(exc)

    summary["offline_boundary"]["socket_restored"] = (
        socket.socket is original_socket
        and socket.create_connection is original_create_connection
        and socket.getaddrinfo is original_getaddrinfo
    )

    case_root = Path(tempfile.mkdtemp(prefix="mfl-installed-offline-qualification-")).resolve()
    records = build_case_fixture(case_root)

    isolated_localappdata = case_root / "localappdata"
    process = start_runtime(
        Path(resolution.executable_path),
        port,
        localappdata=isolated_localappdata,
    )
    network_monitor = RuntimeNetworkMonitor(process.pid)
    network_monitor.start()
    runtime_binding_before = False
    try:
        health = wait_json(f"{base_url}/api/health", timeout_s=180)
        runtime_binding_before = verify_runtime_instance(base_url, process.qa_service_instance)
        if not runtime_binding_before:
            raise RuntimeError("runtime_instance_mismatch_before_actions")
        version = request_json("GET", f"{base_url}/api/version")
        root_page = urllib.request.urlopen(f"{base_url}/", timeout=30).read().decode("utf-8", errors="replace")
        document_status = request_json("GET", f"{base_url}/api/document-intelligence/status")

        summary["feature_results"]["launch"] = {
            "health": health,
            "version": version,
            "rootStatus": 200,
            "rootHasWorkbench": "workbench" in root_page.lower(),
            "processPath": health.get("process_path") or health.get("processPath") or "",
        }
        summary["feature_results"]["document_intelligence_status"] = document_status

        activate = request_json("POST", f"{base_url}/api/activate-corpus", {"case_root": str(case_root)})
        summary["feature_results"]["activate_corpus"] = activate

        request_json("POST", f"{base_url}/api/corpus-rebuild-index", {})
        request_json("POST", f"{base_url}/api/corpus-ocr/start", {"approved": True, "language": "eng"})
        ocr_deadline = time.monotonic() + 180
        corpus_ocr = {}
        while time.monotonic() < ocr_deadline:
            corpus_ocr = request_json("GET", f"{base_url}/api/corpus-ocr/status")
            if corpus_ocr.get("status") not in {"queued", "running"}:
                break
            time.sleep(1)
        summary["feature_results"]["corpus_ocr"] = corpus_ocr
        inventory = request_json("GET", f"{base_url}/api/corpus-inventory")
        retrieval_status = request_json("GET", f"{base_url}/api/retrieval-workbench/status")
        retrieval_search = request_json(
            "POST",
            f"{base_url}/api/retrieval-workbench/search",
            {"query": "changed schools", "include_private_records": True, "include_authority": False, "top_k": 5},
        )
        ask_school = request_json("POST", f"{base_url}/ask", {"question": "changed schools", "search_mode": "my_records"})
        ask_pdf = request_json("POST", f"{base_url}/ask", {"question": "ordinary PDF", "search_mode": "my_records"})
        ask_pii = request_json("POST", f"{base_url}/ask", {"question": "example.com", "search_mode": "my_records"})

        scan_source = case_root / "02_PRIVATE_FORENSIC_MASTER" / "files" / "scan.pdf"
        # These operations must execute in the candidate, never in the driver
        # Python environment. The driver creates fixtures and verifies returned
        # bytes; it does not supply a substitute parser/privacy/OCR service.
        def record_action(record_id: str, action: str, **options: Any) -> dict[str, Any]:
            integrity = request_json("GET", f"{base_url}/api/records/{record_id}/integrity")
            token = integrity["preview"]["token"]
            return request_json("POST", f"{base_url}/api/records/{record_id}/{action}",
                                {"source_token": token, "approved": True, **options})

        pdf_previews = []
        for record_id, page in (("REC-PDF", 1), ("REC-PDF", 2), ("REC-OCR", 1)):
            integrity = request_json("GET", f"{base_url}/api/records/{record_id}/integrity")
            data, receipt = verified_pdf_page(base_url, integrity["preview"], page,
                                               instance=process.qa_service_instance)
            filename = f"fictional-{record_id}-page-{page}.png"
            (evidence_root / filename).write_bytes(data)
            pdf_previews.append({**receipt, "record_id": record_id, "artifact": filename})
        summary["feature_results"]["pdf_raster_previews"] = pdf_previews

        pii_integrity = request_json("GET", f"{base_url}/api/records/REC-PII-TXT/integrity")
        pii_blocks = request_json("GET", f"{base_url}/api/records/REC-PII-TXT/blocks")["blocks"]
        pii_privacy = record_action("REC-PII-TXT", "privacy-scan", run_presidio=True)
        pii_redaction_proposal = record_action("REC-PII-TXT", "redaction-proposal", run_presidio=True)
        pii_redacted_copy = record_action("REC-PII-TXT", "redacted-copy", run_presidio=True,
                                          reviewer="fictional-release-qa")
        redacted_artifact = pii_redacted_copy.get("artifacts", {}).get("redacted_copy", {})
        redacted_bytes = verified_artifact_bytes(base_url, redacted_artifact)
        redacted_receipt = request_json("GET", f"{base_url}/api/artifacts/{redacted_artifact['artifact_id']}/receipt")
        docx_parse_fallback = record_action("REC-DOCX", "parse", run_docling=False, run_presidio=False)
        docx_parse = record_action("REC-DOCX", "parse", run_docling=True, run_presidio=True)
        pdf_parse = record_action("REC-PDF", "parse", run_docling=True, run_presidio=False)
        scan_hash = sha256_file(scan_source)
        ocr_result = record_action("REC-OCR", "ocr", language="eng")
        ocr_artifacts = dict(ocr_result.get("artifacts") or {})
        ocr_pdf_bytes = b""
        ocr_sidecar_bytes = b""
        ocr_text = ""
        if "pdf" in ocr_artifacts and "sidecar" in ocr_artifacts:
            ocr_pdf_artifact = ocr_artifacts["pdf"]
            ocr_sidecar_artifact = ocr_artifacts["sidecar"]
            ocr_pdf_bytes = verified_artifact_bytes(base_url, ocr_pdf_artifact)
            ocr_sidecar_bytes = verified_artifact_bytes(base_url, ocr_sidecar_artifact)
            ocr_text = PdfReader(io.BytesIO(ocr_pdf_bytes)).pages[0].extract_text() or ""
        ocr_comparison = ocr_result.get("comparison", {})
        duplicates = request_json("GET", f"{base_url}/api/records/REC-DUP-A/duplicates")
        compare = request_json("POST", f"{base_url}/api/records/compare",
                               {"left_record_id": "REC-DUP-A", "right_record_id": "REC-DUP-B"})
        changed_compare = request_json("POST", f"{base_url}/api/records/compare",
                                       {"left_record_id": "REC-DUP-A", "right_record_id": "REC-CHANGED"})

        summary["feature_results"].update(
            {
                "inventory": inventory,
                "retrieval_status": retrieval_status,
                "retrieval_search": retrieval_search,
                "ask": {
                    "school_change": {
                        "grounded": ask_school.get("grounded"),
                        "failure_class": ask_school.get("failure_class"),
                        "source_card_count": ask_school.get("source_card_count"),
                        "citations": ask_school.get("citations", [])[:3],
                    },
                    "ordinary_pdf": {
                        "grounded": ask_pdf.get("grounded"),
                        "failure_class": ask_pdf.get("failure_class"),
                        "source_card_count": ask_pdf.get("source_card_count"),
                        "citations": ask_pdf.get("citations", [])[:3],
                    },
                    "pii_note": {
                        "grounded": ask_pii.get("grounded"),
                        "failure_class": ask_pii.get("failure_class"),
                        "source_card_count": ask_pii.get("source_card_count"),
                        "citations": ask_pii.get("citations", [])[:3],
                    },
                },
                "document_intelligence": {
                    "integrity": pii_integrity,
                    "privacy_scan": pii_privacy,
                    "blocks": pii_blocks,
                    "redaction_proposal": pii_redaction_proposal,
                    "redacted_copy": pii_redacted_copy,
                    "fallback_parse": docx_parse_fallback,
                    "docling_parse": docx_parse,
                    "pdf_parse": pdf_parse,
                    "ocr": {
                        "result": ocr_result,
                        "pdf_hash": hashlib.sha256(ocr_pdf_bytes).hexdigest() if ocr_pdf_bytes else None,
                        "sidecar_hash": hashlib.sha256(ocr_sidecar_bytes).hexdigest() if ocr_sidecar_bytes else None,
                        "pdf_text": ocr_text[:2000],
                        "comparison": ocr_comparison,
                    },
                    "redaction": pii_redacted_copy,
                    "duplicates": duplicates,
                    "compare": compare,
                    "changed_copy": changed_compare,
                },
            }
        )

        summary["inventory_result"] = {
            "status": inventory.get("status"),
            "records": inventory.get("records"),
            "searchable_records": inventory.get("searchable_records"),
            "ocr_candidates": inventory.get("ocr_candidates"),
        }

        fail_conditions = {
            "pdf_raster_pages": len(pdf_previews) != 3 or any(row["status"] != "pass" for row in pdf_previews),
            "pdf_distinct_pages": pdf_previews[0]["sha256"] == pdf_previews[1]["sha256"],
            "pdf_page_count": pdf_previews[0]["page_count"] != 2 or pdf_previews[1]["page_count"] != 2,
            "corpus_ocr": corpus_ocr.get("status") != "completed"
            or corpus_ocr.get("failed") != 0 or int(corpus_ocr.get("completed") or 0) < 1,
            "runtime_health": health.get("status") not in {"ok", "ready", "degraded"},
            "workbench_root": not summary["feature_results"]["launch"]["rootHasWorkbench"],
            "retrieval_status": retrieval_status.get("status") == "blocked",
            "retrieval_search": retrieval_search.get("status") not in {"pass", "no_matches"},
            "grounded_school_answer": ask_school.get("source_card_count", 0) <= 0,
            "grounded_private_answer": ask_pii.get("source_card_count", 0) <= 0,
            "no_automatic_install": document_status.get("automatic_install") is not False,
            "no_document_network": document_status.get("network_used") is not False,
            "presidio_available": not any(
                row.get("available")
                for row in document_status.get("adapters", [])
                if row.get("adapter_id") == "presidio"
            ),
            "deterministic_fallback": docx_parse_fallback.get("selected_extractor") != "deterministic_baseline",
            "fallback_reason": "deterministic_baseline_selected"
            not in str(docx_parse_fallback.get("selection_reason") or ""),
            "docling_or_fallback": docx_parse.get("selected_extractor") not in {"docling", "deterministic_baseline"},
            "docling_engine": docx_parse.get("selected_extractor") != "docling",
            "privacy_worker": pii_privacy.get("privacy_review", {}).get("presidio_status") != "pass",
            "redacted_copy": pii_redacted_copy.get("status") != "pass"
            or pii_redacted_copy.get("review_required") is not True
            or not redacted_bytes or b"jane.example@example.com" in redacted_bytes,
            "redaction_receipt": redacted_receipt.get("artifact_type") != "redacted_copy",
            "ocr_status": not ocr_completed(ocr_result, ocr_text, ocr_sidecar_bytes),
            "original_immutable": ocr_result.get("original_modified") is not False,
            "original_hash_unchanged": sha256_file(scan_source) != scan_hash,
            "duplicate_detection": duplicates.get("exact_duplicate") is not True,
            "record_comparison": compare.get("exact_duplicate") is not True,
            "changed_copy": changed_compare.get("exact_duplicate") is not False
            or not changed_compare.get("field_differences", {}).get("source_hash"),
        }
        failed_checks = [name for name, failed in fail_conditions.items() if failed]
        summary["qualification_checks"] = {
            name: {"status": "fail" if failed else "pass"} for name, failed in fail_conditions.items()
        }
        if failed_checks:
            summary["failed_qualification_checks"] = failed_checks
            summary["blockers"].append("one_or_more_feature_checks_failed")

        summary["feature_check_status"] = "pass" if not summary["blockers"] else "blocked"
        if args.hold_seconds:
            (evidence_root / "browser-ready.json").write_text(
                json.dumps({"base_url": base_url, "fictional_only": True,
                            "runtime_sha256": summary["runtime_sha256"],
                            "feature_check_status": summary["feature_check_status"]}), encoding="utf-8")
            deadline = time.monotonic() + args.hold_seconds
            while time.monotonic() < deadline and not (evidence_root / "stop-probe").exists():
                time.sleep(1)
    except Exception as exc:
        summary["blockers"].append(f"qualification_exception:{type(exc).__name__}")
        summary["feature_check_status"] = "blocked"
    finally:
        runtime_network = network_monitor.stop()
        summary["offline_boundary"]["runtime_network_observation"] = runtime_network
        if runtime_network.get("external_connection_count"):
            summary["blockers"].append("installed_runtime_external_network_connection_observed")
            summary["feature_check_status"] = "blocked"
        if runtime_network.get("errors") or not runtime_network.get("sample_count"):
            summary["blockers"].append("runtime_network_observation_incomplete")
        summary["request_events"] = list(REQUEST_EVENTS)
        instance = str(getattr(process, "qa_service_instance", ""))
        summary["runtime_instance_verified"] = runtime_binding_before and verify_runtime_instance(base_url, instance)
        summary["runtime_instance_proof"] = "fresh secret instance header on health before and after actions on the same loopback origin"
        if not summary["runtime_instance_verified"]:
            summary["blockers"].append("runtime_instance_not_verified")
            summary["feature_check_status"] = "blocked"
        try:
            process.terminate()
            process.wait(timeout=30)
        except Exception:
            process.kill()

    # TCP polling and dead proxies do not prove zero UDP/DNS/native requests.
    # A successful explicit frozen-runtime run is not an installed-package run.
    summary["feature_blockers"] = list(summary["blockers"])
    summary["blockers"].append("os_level_zero_network_proof_not_executed")
    if resolution.source != "appx_package":
        summary["blockers"].append("installed_msix_not_tested")
    summary["qualification_status"] = "blocked"
    summary_path = evidence_root / "installed-offline-qualification.json"
    text_path = evidence_root / "installed-offline-qualification.txt"
    summary["files"] = {
        "installed_offline_qualification_json": str(summary_path),
        "installed_offline_qualification_txt": str(text_path),
        "bundled_engine_inventory_json": str(evidence_root / "bundled-engine-inventory.json"),
    }
    summary_path.write_text(json.dumps(summary, indent=2, sort_keys=True), encoding="utf-8")
    text_path.write_text(
        "\n".join(
            [
                f"Runtime source: {summary['runtime_resolution']['source']}",
                "Method: canonical frozen HTTP calls; driver-only socket guard and best-effort TCP observation, not OS network proof",
                f"Feature result: {summary['qualification_status']}",
                f"Scoped API result: {summary['feature_check_status']}",
                f"Network attempt: {summary['offline_boundary'].get('external_request')}",
                f"Network restoration: {summary['offline_boundary'].get('socket_restored')}",
                f"Inventory records: {summary['inventory_result'].get('records')}",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    print(json.dumps({"report": str(summary_path), "feature_check_status": summary["feature_check_status"],
                      "qualification_status": summary["qualification_status"],
                      "blockers": summary["blockers"]}, indent=2, sort_keys=True))
    return 0 if summary["qualification_status"] == "pass" else 2


if __name__ == "__main__":
    raise SystemExit(main())
