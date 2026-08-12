from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import shutil
import sys
from pathlib import Path

from docx import Document
from fpdf import FPDF
from PIL import Image, ImageDraw

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT / "src") not in sys.path:
    sys.path.insert(0, str(ROOT / "src"))

from maine_family_law_llm.case_corpus_builder import build_case_corpus


FIXED_MTIME = 1_767_225_600  # 2026-01-01T00:00:00Z


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def write_text(path: Path, text: str) -> None:
    path.write_text(text.rstrip() + "\n", encoding="utf-8")
    os.utime(path, (FIXED_MTIME, FIXED_MTIME))


def write_docx(path: Path, *, changed: bool = False) -> None:
    document = Document()
    document.add_heading("FICTIONAL DISCOVERY REQUEST", level=1)
    document.add_paragraph("DEMONSTRATION MATTER — NO REAL PERSON OR FACT")
    document.add_paragraph("Request 1: Produce the fictional school calendar referenced in the communication record.")
    document.add_paragraph("Request 2: Produce the fictional attachment called schedule-demo.pdf.")
    if changed:
        document.add_paragraph("Changed-copy note: Request 3 asks for a fictional transportation log.")
    document.save(path)
    os.utime(path, (FIXED_MTIME, FIXED_MTIME))


def write_scan(path: Path) -> None:
    image_path = path.with_suffix(".png")
    image = Image.new("RGB", (1240, 1754), "white")
    draw = ImageDraw.Draw(image)
    draw.text((80, 100), "FICTIONAL SCANNED NOTICE — OCR REQUIRED", fill="black")
    draw.text((80, 165), "Demonstration hearing date: February 20, 2026 at 09:00", fill="black")
    draw.text((80, 230), "No real person, court filing, docket, or fact.", fill="black")
    image.save(image_path)
    pdf = FPDF(unit="pt", format="letter")
    pdf.add_page()
    pdf.image(str(image_path), x=36, y=36, w=540)
    pdf.output(str(path))
    image_path.unlink()
    os.utime(path, (FIXED_MTIME, FIXED_MTIME))


def build_fixture(run_root: Path) -> dict[str, object]:
    input_root = run_root / "inputs"
    output_root = run_root / "built"
    input_root.mkdir(parents=True, exist_ok=True)
    output_root.mkdir(parents=True, exist_ok=True)

    write_text(
        input_root / "01_fictional_complaint.txt",
        """FICTIONAL COMPLAINT — DEMONSTRATION ONLY
Parent Alpha and Parent Beta are fictional labels. Child Gamma is a fictional label.
The pleading requests a review of parental rights and responsibilities.
Valid citation fixture: 19-A M.R.S. § 1653.
Every claim requires human review and exact-source verification.""",
    )
    write_text(
        input_root / "02_fictional_motion.txt",
        """FICTIONAL MOTION — DEMONSTRATION ONLY
The motion alleges that an exchange occurred on January 15, 2026.
The communication record instead reports January 16, 2026.
This conflict is intentional and must not be resolved automatically.""",
    )
    write_text(
        input_root / "03_initial_order.txt",
        """FICTIONAL INITIAL ORDER — DEMONSTRATION ONLY
Entered January 5, 2026. Exact term: Exchanges occur Fridays at 17:00 at Demo Library.
This term is a source-bound candidate and requires reviewer confirmation.""",
    )
    write_text(
        input_root / "04_modifying_order.txt",
        """FICTIONAL AMENDED ORDER — DEMONSTRATION ONLY
Entered January 12, 2026. This fictional order modifies the exchange term in the January 5 order.
Exact amended term: Exchanges occur Saturdays at 10:00 at Demo Community Center.
The application must not decide operative status without review.""",
    )
    write_scan(input_root / "05_scanned_hearing_notice.pdf")

    request_docx = input_root / "06_discovery_request.docx"
    write_docx(request_docx)
    duplicate_docx = input_root / "07_discovery_request_exact_duplicate.docx"
    shutil.copyfile(request_docx, duplicate_docx)
    os.utime(duplicate_docx, (FIXED_MTIME, FIXED_MTIME))
    write_docx(input_root / "08_discovery_request_changed_copy.docx", changed=True)

    write_text(
        input_root / "09_communication_missing_attachment.eml",
        """From: fictional.alpha.invalid
To: fictional.beta.invalid
Date: Fri, 16 Jan 2026 10:00:00 -0500
Subject: FICTIONAL exchange and missing attachment
Message-ID: <fictional-ga-001.invalid>
MIME-Version: 1.0
Content-Type: text/plain; charset=utf-8

DEMONSTRATION ONLY. The fictional exchange occurred January 16, 2026.
Please review the attached schedule-demo.pdf. The attachment is intentionally missing.""",
    )

    with (input_root / "10_fictional_docket.csv").open("w", newline="", encoding="utf-8") as stream:
        writer = csv.writer(stream)
        writer.writerow(["sequence", "date", "description", "source"])
        writer.writerow(["1", "2026-01-05", "FICTIONAL initial order entered", "03_initial_order.txt"])
        writer.writerow(["2", "2026-01-12", "FICTIONAL modifying order entered", "04_modifying_order.txt"])
        writer.writerow(["3", "2026-02-01", "FICTIONAL notice references missing attachment", "schedule-demo.pdf"])
    os.utime(input_root / "10_fictional_docket.csv", (FIXED_MTIME, FIXED_MTIME))

    write_text(
        input_root / "11_partial_discovery_response.txt",
        """FICTIONAL PARTIAL DISCOVERY RESPONSE — DEMONSTRATION ONLY
Response to Request 1: A fictional calendar excerpt is produced.
Response to Request 2: No schedule-demo.pdf is included. Production is incomplete.
Privilege, completeness, and compliance remain review-required.""",
    )
    write_text(
        input_root / "12_draft_with_unsupported_claims.txt",
        """FICTIONAL DRAFT — NOT FOR FILING
Unsupported claim: The court always selects Parent Alpha.
Unsupported claim: The missing schedule-demo.pdf proves the exchange date.
Qualification: No source in this fixture supports either proposition.
Missing context: the scanned notice has not been OCR-verified.""",
    )
    write_text(
        input_root / "13_stale_form_fixture.txt",
        """FICTIONAL STALE FORM FIXTURE — DO NOT FILE
Form ID: DEMO-FM-999
Revision date: 2019-01
Currentness: intentionally stale and unverified.
The application must warn the user and must not represent this as a current court form.""",
    )
    write_text(
        input_root / "14_citation_and_quote_fixtures.txt",
        """FICTIONAL VERIFICATION FIXTURES
Valid Maine citation fixture: 19-A M.R.S. § 1653.
Fake citation fixture: 19-Z M.R.S. § 9999.
Exact quote fixture: "Exchanges occur Saturdays at 10:00 at Demo Community Center."
Mismatched quote fixture: "Exchanges occur Sundays at noon at Demo Courthouse."
The mismatched quote does not appear in either fictional order.""",
    )

    result = build_case_corpus(
        repo_root=ROOT,
        source_roots=[input_root],
        output_root=output_root,
        case_name="Fictional GA Matter 2026",
    )
    files = sorted(path for path in input_root.iterdir() if path.is_file())
    manifest = {
        "schema_version": "fictional_ga_matter_v1",
        "fictional": True,
        "private_or_real_data": False,
        "matter_label": "Fictional GA Matter 2026",
        "case_root": str(result.case_root),
        "record_count": len(files),
        "records": [
            {
                "filename": path.name,
                "sha256": sha256(path),
                "bytes": path.stat().st_size,
            }
            for path in files
        ],
        "required_fixtures": {
            "pleadings": ["01_fictional_complaint.txt", "02_fictional_motion.txt"],
            "orders": ["03_initial_order.txt", "04_modifying_order.txt"],
            "scanned_pdf": "05_scanned_hearing_notice.pdf",
            "docx": "06_discovery_request.docx",
            "exact_duplicate": "07_discovery_request_exact_duplicate.docx",
            "changed_copy": "08_discovery_request_changed_copy.docx",
            "communication_missing_attachment": "09_communication_missing_attachment.eml",
            "docket": "10_fictional_docket.csv",
            "partial_discovery_response": "11_partial_discovery_response.txt",
            "unsupported_draft": "12_draft_with_unsupported_claims.txt",
            "stale_form": "13_stale_form_fixture.txt",
            "citation_and_quote_fixtures": "14_citation_and_quote_fixtures.txt"
        },
    }
    manifest_path = run_root / "fictional_ga_matter_manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2, sort_keys=True) + "\n", encoding="utf-8")
    return {**manifest, "manifest_path": str(manifest_path)}


def main() -> int:
    parser = argparse.ArgumentParser(description="Build the deterministic fictional GA E2E matter.")
    parser.add_argument(
        "--run-root",
        default=str(ROOT / "dist" / "ga_today" / "e2e_runtime" / "fictional_ga_matter_20260811"),
    )
    args = parser.parse_args()
    run_root = Path(args.run_root).resolve()
    result = build_fixture(run_root)
    print(json.dumps(result, indent=2, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
