from __future__ import annotations

import hashlib
import io
import re
from typing import Any

from docx import Document
from pypdf import PdfReader

from .contracts import DocumentBlock

MAX_BLOCKS = 20_000
MAX_TEXT_CHARS = 8_000_000
_HEADING_RE = re.compile(r"^(?:[A-Z][A-Z0-9 &'()\-.,]{3,}|\d+(?:\.\d+)*\.?\s+\S.+)$")
_NUMBERED_RE = re.compile(r"^\s*(?:\d+\.|\([a-z0-9]+\)|[A-Z]\.)\s+")
_SIGNATURE_RE = re.compile(r"(?i)^(?:dated:|respectfully submitted|signature|by:|attorney for|pro se)")


def _block_id(source_hash: str, order: int, kind: str, text: str) -> str:
    material = f"{source_hash}:{order}:{kind}:{text}".encode("utf-8", errors="replace")
    return f"blk_{hashlib.sha256(material).hexdigest()[:20]}"


def _kind_for_line(line: str) -> str:
    stripped = line.strip()
    if not stripped:
        return "blank"
    if _SIGNATURE_RE.match(stripped):
        return "signature"
    if "\t" in line or re.search(r"\s{3,}", line):
        return "table_row"
    if _HEADING_RE.match(stripped) and len(stripped) <= 180:
        return "heading"
    if _NUMBERED_RE.match(stripped):
        return "numbered_paragraph"
    return "paragraph"


def _append_line_blocks(
    *, source_hash: str, text: str, page_number: int, blocks: list[DocumentBlock], cursor: int
) -> int:
    for raw_line in text.splitlines():
        if len(blocks) >= MAX_BLOCKS:
            break
        clean = raw_line.strip()
        if not clean:
            cursor += len(raw_line) + 1
            continue
        start = cursor + max(0, raw_line.find(clean))
        end = start + len(clean)
        kind = _kind_for_line(raw_line)
        order = len(blocks) + 1
        blocks.append(
            DocumentBlock(
                block_id=_block_id(source_hash, order, kind, clean),
                kind=kind,
                text=clean[:20_000],
                page_number=page_number,
                order=order,
                char_start=start,
                char_end=end,
                metadata={"extractor": "deterministic_baseline"},
            )
        )
        cursor += len(raw_line) + 1
    return cursor


def extract_baseline_blocks(data: bytes, suffix: str, source_hash: str) -> dict[str, Any]:
    suffix = suffix.lower()
    blocks: list[DocumentBlock] = []
    warnings: list[str] = []
    cursor = 0
    page_count = 0

    if suffix == ".pdf":
        try:
            reader = PdfReader(io.BytesIO(data))
            page_count = len(reader.pages)
            for page_number, page in enumerate(reader.pages, start=1):
                try:
                    text = page.extract_text() or ""
                except Exception as exc:
                    text = ""
                    warnings.append(f"page_{page_number}_native_text_error:{exc.__class__.__name__}")
                cursor = _append_line_blocks(
                    source_hash=source_hash,
                    text=text[:MAX_TEXT_CHARS],
                    page_number=page_number,
                    blocks=blocks,
                    cursor=cursor,
                )
                if len(blocks) >= MAX_BLOCKS:
                    warnings.append("baseline_block_limit_reached")
                    break
        except Exception as exc:
            warnings.append(f"baseline_pdf_error:{exc.__class__.__name__}")
    elif suffix == ".docx":
        try:
            document = Document(io.BytesIO(data))
            for paragraph in document.paragraphs:
                cursor = _append_line_blocks(
                    source_hash=source_hash,
                    text=paragraph.text,
                    page_number=0,
                    blocks=blocks,
                    cursor=cursor,
                )
            for table_index, table in enumerate(document.tables, start=1):
                for row_index, row in enumerate(table.rows, start=1):
                    text = " | ".join(cell.text.strip() for cell in row.cells)
                    if not text.strip(" |"):
                        continue
                    order = len(blocks) + 1
                    blocks.append(
                        DocumentBlock(
                            block_id=_block_id(source_hash, order, "table_row", text),
                            kind="table_row",
                            text=text[:20_000],
                            order=order,
                            char_start=cursor,
                            char_end=cursor + len(text),
                            metadata={
                                "extractor": "deterministic_baseline",
                                "table_index": table_index,
                                "row_index": row_index,
                            },
                        )
                    )
                    cursor += len(text) + 1
                    if len(blocks) >= MAX_BLOCKS:
                        warnings.append("baseline_block_limit_reached")
                        break
        except Exception as exc:
            warnings.append(f"baseline_docx_error:{exc.__class__.__name__}")
    else:
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("latin-1", errors="replace")
        cursor = _append_line_blocks(
            source_hash=source_hash,
            text=text[:MAX_TEXT_CHARS],
            page_number=0,
            blocks=blocks,
            cursor=cursor,
        )

    kinds: dict[str, int] = {}
    for block in blocks:
        kinds[block.kind] = kinds.get(block.kind, 0) + 1
    return {
        "schema_version": "document_intelligence_baseline_v1",
        "status": "pass" if blocks else "review_required",
        "source_sha256": source_hash,
        "page_count": page_count,
        "block_count": len(blocks),
        "block_kind_counts": kinds,
        "blocks": [block.as_dict() for block in blocks],
        "warnings": warnings,
        "review_required": True,
    }
