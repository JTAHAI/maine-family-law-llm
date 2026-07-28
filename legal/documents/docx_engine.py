"""Safe adapters for creating and editing Word documents.

Tracked-change editing uses the MIT-licensed ``docx-editor`` package by Pablo
Speciale when installed. The adapter deliberately exposes only a bounded subset
of its API and always writes a new artifact; it never overwrites an imported
source document.
"""

from __future__ import annotations

import hashlib
import importlib
import importlib.metadata
import os
import re
import secrets
import shutil
from pathlib import Path
from typing import Any, Iterable

from .workspace import DocumentWorkspaceError, _atomic_write

MAX_DOCX_BYTES = 50 * 1024 * 1024
MAX_EDIT_OPERATIONS = 100
MAX_EDIT_TEXT_CHARS = 50_000
_PARAGRAPH_REF_RE = re.compile(r"^P[1-9][0-9]*#[a-f0-9]{4,64}$")
_ALLOWED_ACTIONS = {"replace", "delete", "insert_after", "rewrite_paragraph", "add_comment"}


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _inside(path: Path, root: Path) -> Path:
    if path.is_symlink():
        raise DocumentWorkspaceError("docx_symlink_refused", "A DOCX symlink was refused.", status_code=409)
    try:
        resolved = path.resolve(strict=True)
        resolved_root = root.resolve(strict=True)
        resolved.relative_to(resolved_root)
    except (FileNotFoundError, ValueError, OSError) as exc:
        raise DocumentWorkspaceError("docx_path_outside_workspace", "The DOCX path is outside the local workspace.", status_code=409) from exc
    return resolved


def engine_status() -> dict[str, Any]:
    status: dict[str, Any] = {
        "engine": "docx-editor",
        "license": "MIT",
        "copyright": "Copyright (c) 2026 Pablo Speciale",
        "tracked_changes_available": False,
        "document_creation_available": False,
        "local_only": True,
        "source_overwrite_allowed": False,
    }
    try:
        importlib.import_module("docx_editor")
        try:
            version = importlib.metadata.version("docx-editor")
        except importlib.metadata.PackageNotFoundError:
            version = "source-tree"
        status.update({"tracked_changes_available": True, "version": version})
    except Exception as exc:
        status["tracked_changes_reason"] = f"docx-editor unavailable: {type(exc).__name__}"
    try:
        importlib.import_module("docx")
        status["document_creation_available"] = True
    except Exception as exc:
        status["document_creation_reason"] = f"python-docx unavailable: {type(exc).__name__}"
    return status


def create_docx_from_text(
    *,
    title: str,
    content: str,
    output_path: Path,
    allowed_output_root: Path,
) -> dict[str, Any]:
    """Create a new review-required DOCX without touching any source document."""
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:
        raise DocumentWorkspaceError("python_docx_unavailable", "Word export requires python-docx.", status_code=503) from exc

    output_path = Path(output_path)
    allowed_output_root = Path(allowed_output_root).resolve(strict=True)
    if output_path.exists() and output_path.is_symlink():
        raise DocumentWorkspaceError("docx_symlink_refused", "A DOCX symlink was refused.", status_code=409)
    resolved_parent = output_path.parent.resolve(strict=True)
    if resolved_parent != allowed_output_root and allowed_output_root not in resolved_parent.parents:
        raise DocumentWorkspaceError("docx_path_outside_workspace", "The DOCX output path is outside the local workspace.", status_code=409)
    if output_path.suffix.lower() != ".docx":
        raise DocumentWorkspaceError("invalid_docx_extension", "Word exports must use the .docx extension.")

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    document.add_heading(str(title)[:240], level=0)
    document.add_paragraph("REVIEW REQUIRED — WORKING DRAFT — NOT FILING-READY")
    for raw_line in str(content or "").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            document.add_paragraph()
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            document.add_heading(stripped[2:], level=1)
        elif re.match(r"^[-*]\s+", stripped):
            document.add_paragraph(re.sub(r"^[-*]\s+", "", stripped), style="List Bullet")
        elif re.match(r"^\d+[.)]\s+", stripped):
            document.add_paragraph(re.sub(r"^\d+[.)]\s+", "", stripped), style="List Number")
        else:
            document.add_paragraph(line)

    temporary = output_path.parent / f".{output_path.name}.{secrets.token_hex(8)}.tmp.docx"
    try:
        document.save(temporary)
        data = temporary.read_bytes()
        if len(data) > MAX_DOCX_BYTES:
            raise DocumentWorkspaceError("docx_export_too_large", "The generated DOCX exceeds the safety limit.", status_code=413)
        _atomic_write(output_path, data)
    finally:
        if temporary.exists():
            temporary.unlink(missing_ok=True)
    return {
        "path": output_path,
        "size_bytes": output_path.stat().st_size,
        "sha256": _sha256(output_path),
        "format": "docx",
        "review_required": True,
        "filing_ready": False,
        "original_preserved": True,
    }


def _normalize_operations(operations: Iterable[dict[str, Any]]) -> list[dict[str, Any]]:
    rows = list(operations or [])
    if not rows:
        raise DocumentWorkspaceError("docx_edits_required", "At least one DOCX edit is required.")
    if len(rows) > MAX_EDIT_OPERATIONS:
        raise DocumentWorkspaceError("too_many_docx_edits", f"No more than {MAX_EDIT_OPERATIONS} edits may be applied at once.", status_code=413)
    normalized: list[dict[str, Any]] = []
    for index, raw in enumerate(rows):
        if not isinstance(raw, dict):
            raise DocumentWorkspaceError("invalid_docx_edit", f"DOCX edit {index + 1} is invalid.")
        action = str(raw.get("action") or "").strip().lower()
        if action not in _ALLOWED_ACTIONS:
            raise DocumentWorkspaceError("invalid_docx_edit_action", f"DOCX edit action '{action}' is not allowed.")
        paragraph = str(raw.get("paragraph") or "").strip()
        if action != "add_comment" and not _PARAGRAPH_REF_RE.fullmatch(paragraph):
            raise DocumentWorkspaceError("invalid_paragraph_reference", "DOCX edits require a current hash-anchored paragraph reference.")
        row: dict[str, Any] = {"action": action}
        if paragraph:
            row["paragraph"] = paragraph
        for key in ("find", "replace_with", "text", "comment"):
            if key in raw:
                value = str(raw.get(key) or "").replace("\x00", "")
                if len(value) > MAX_EDIT_TEXT_CHARS:
                    raise DocumentWorkspaceError("docx_edit_text_too_large", f"DOCX edit field '{key}' is too large.", status_code=413)
                row[key] = value
        occurrence = raw.get("occurrence", 0)
        try:
            occurrence = int(occurrence)
        except (TypeError, ValueError) as exc:
            raise DocumentWorkspaceError("invalid_docx_occurrence", "DOCX occurrence must be an integer.") from exc
        if occurrence < 0 or occurrence > 100_000:
            raise DocumentWorkspaceError("invalid_docx_occurrence", "DOCX occurrence is outside the allowed range.")
        row["occurrence"] = occurrence
        normalized.append(row)
    return normalized


def list_docx_paragraphs(
    *,
    source_path: Path,
    allowed_source_root: Path,
    start: int = 1,
    limit: int = 200,
) -> dict[str, Any]:
    """Return bounded, hash-anchored paragraph references from a preserved DOCX."""
    source = _inside(Path(source_path), Path(allowed_source_root))
    if source.suffix.lower() != ".docx" or source.stat().st_size > MAX_DOCX_BYTES:
        raise DocumentWorkspaceError("invalid_docx_source", "The preserved source is not an allowed DOCX.", status_code=409)
    try:
        from docx_editor import Document
    except Exception as exc:
        raise DocumentWorkspaceError("docx_editor_unavailable", "Tracked Word editing requires the MIT-licensed docx-editor package.", status_code=503) from exc
    start = max(1, min(int(start), 100_000))
    limit = max(1, min(int(limit), 500))
    workspace_dir = source.parent / ".docx-editor-workspaces"
    workspace_dir.mkdir(mode=0o700, exist_ok=True)
    with Document.open(source, author="Maine Family Law LLM User", force_recreate=True, workspace_dir=workspace_dir) as document:
        paragraphs = document.list_paragraphs_structured(start=start, limit=limit)
        total = document.paragraph_count()
        rows = [{"index": int(item.index), "ref": str(item.ref), "text": str(item.text)[:10_000]} for item in paragraphs]
    return {"paragraphs": rows, "total": total, "start": start, "limit": limit, "source_sha256": _sha256(source)}


def tracked_edit_copy(
    *,
    source_path: Path,
    allowed_source_root: Path,
    output_path: Path,
    allowed_output_root: Path,
    operations: Iterable[dict[str, Any]],
    author: str = "Maine Family Law LLM User",
) -> dict[str, Any]:
    """Apply bounded tracked edits to a new copy of a preserved DOCX source."""
    source = _inside(Path(source_path), Path(allowed_source_root))
    if source.suffix.lower() != ".docx" or source.stat().st_size > MAX_DOCX_BYTES:
        raise DocumentWorkspaceError("invalid_docx_source", "The preserved source is not an allowed DOCX.", status_code=409)
    output_path = Path(output_path)
    output_root = Path(allowed_output_root).resolve(strict=True)
    output_parent = output_path.parent.resolve(strict=True)
    if output_parent != output_root and output_root not in output_parent.parents:
        raise DocumentWorkspaceError("docx_path_outside_workspace", "The DOCX output path is outside the local workspace.", status_code=409)
    if output_path.suffix.lower() != ".docx" or output_path.exists():
        raise DocumentWorkspaceError("docx_output_conflict", "Tracked edits require a new .docx output path.", status_code=409)
    normalized = _normalize_operations(operations)
    try:
        from docx_editor import Document
    except Exception as exc:
        raise DocumentWorkspaceError("docx_editor_unavailable", "Tracked Word editing requires the MIT-licensed docx-editor package.", status_code=503) from exc

    temporary = output_path.parent / f".{output_path.name}.{secrets.token_hex(8)}.working.docx"
    shutil.copyfile(source, temporary)
    try:
        workspace_dir = output_path.parent / ".docx-editor-workspaces"
        workspace_dir.mkdir(mode=0o700, exist_ok=True)
        edit_results: list[dict[str, Any]] = []
        with Document.open(temporary, author=(str(author).strip()[:120] or "Reviewer"), force_recreate=True, workspace_dir=workspace_dir) as document:
            for row in normalized:
                action = row["action"]
                if action == "replace":
                    new_ref = document.replace(row.get("find", ""), row.get("replace_with", ""), paragraph=row["paragraph"], occurrence=row["occurrence"])
                    edit_results.append({"action": action, "paragraph": row["paragraph"], "new_ref": str(new_ref)})
                elif action == "delete":
                    new_ref = document.delete(row.get("text", row.get("find", "")), paragraph=row["paragraph"], occurrence=row["occurrence"])
                    edit_results.append({"action": action, "paragraph": row["paragraph"], "new_ref": str(new_ref)})
                elif action == "insert_after":
                    new_ref = document.insert_after(row.get("find", ""), row.get("text", ""), paragraph=row["paragraph"], occurrence=row["occurrence"])
                    edit_results.append({"action": action, "paragraph": row["paragraph"], "new_ref": str(new_ref)})
                elif action == "rewrite_paragraph":
                    result = document.rewrite_paragraph(row["paragraph"], row.get("text", ""))
                    edit_results.append({"action": action, "paragraph": row["paragraph"], "new_ref": str(result)})
                elif action == "add_comment":
                    comment = document.add_comment(row.get("find", row.get("text", "")), row.get("comment", "Review required"), occurrence=row["occurrence"])
                    edit_results.append({"action": action, "comment_id": int(getattr(comment, "id", -1))})
            document.save()
        data = temporary.read_bytes()
        if len(data) > MAX_DOCX_BYTES:
            raise DocumentWorkspaceError("docx_export_too_large", "The edited DOCX exceeds the safety limit.", status_code=413)
        _atomic_write(output_path, data)
    finally:
        temporary.unlink(missing_ok=True)
    return {
        "path": output_path,
        "size_bytes": output_path.stat().st_size,
        "sha256": _sha256(output_path),
        "source_sha256": _sha256(source),
        "edit_count": len(normalized),
        "edits": edit_results,
        "tracked_changes": True,
        "review_required": True,
        "filing_ready": False,
        "original_preserved": True,
    }
