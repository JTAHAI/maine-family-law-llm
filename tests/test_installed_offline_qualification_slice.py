from __future__ import annotations

import json
import socket
import subprocess
import time
import urllib.request
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document

from legal.document_intelligence.service import analyze_document, document_intelligence_status
from legal.retrieval.optional_backends import optional_backend_status
from maine_family_law_llm.installed_runtime import resolve_installed_runtime_executable
from maine_family_law_llm.local_only_boundary import LocalOnlyNetworkBlocked, local_only_network_boundary


def _make_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("Motion for Temporary Relief", level=1)
    doc.add_paragraph("The child changed schools on January 3, 2026.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Date"
    table.cell(0, 1).text = "Event"
    table.cell(1, 0).text = "2026-01-03"
    table.cell(1, 1).text = "School change"
    doc.save(path)


def _free_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        sock.bind(("127.0.0.1", 0))
        return int(sock.getsockname()[1])


def _wait_for_json(url: str, timeout_s: int = 120) -> dict[str, object]:
    deadline = time.time() + timeout_s
    last_error = "not_started"
    while time.time() < deadline:
        try:
            with urllib.request.urlopen(url, timeout=5) as response:
                return json.loads(response.read().decode("utf-8"))
        except Exception as exc:  # noqa: BLE001
            last_error = f"{exc.__class__.__name__}: {exc}"
            time.sleep(1.5)
    raise TimeoutError(f"Timed out waiting for {url}: {last_error}")


def test_resolve_installed_runtime_executable_prefers_appx_location(monkeypatch, tmp_path: Path) -> None:
    payload = {
        "PackageFullName": "TAHAIWebServices.MaineFamilyLawLLM_6.0.4.0_x64__k9af96g77tmj4",
        "InstallLocation": str(tmp_path / "installed"),
        "Version": "6.0.4.0",
    }
    install = Path(payload["InstallLocation"])
    install.mkdir()
    (install / "MaineFamilyLawLLM.exe").write_text("stub", encoding="utf-8")

    completed = subprocess.CompletedProcess(
        args=["powershell"],
        returncode=0,
        stdout=json.dumps(payload),
        stderr="",
    )
    monkeypatch.setattr(subprocess, "run", lambda *args, **kwargs: completed)

    resolution = resolve_installed_runtime_executable()
    assert resolution.available is True
    assert resolution.source == "appx_package"
    assert resolution.executable_path == str(install / "MaineFamilyLawLLM.exe")
    assert resolution.install_location == str(install)


def test_local_only_network_boundary_blocks_and_restores_socket_hooks() -> None:
    original_socket = socket.socket
    original_create_connection = socket.create_connection
    original_getaddrinfo = socket.getaddrinfo
    with local_only_network_boundary():
        with pytest.raises(LocalOnlyNetworkBlocked):
            urllib.request.urlopen("https://example.com", timeout=5)
    assert socket.socket is original_socket
    assert socket.create_connection is original_create_connection
    assert socket.getaddrinfo is original_getaddrinfo


def test_document_intelligence_fallback_status_is_honest(tmp_path: Path) -> None:
    case_root = tmp_path / "case"
    case_root.mkdir()
    source = case_root / "motion.docx"
    _make_docx(source)
    before = source.read_bytes()

    result = analyze_document(case_root=case_root, source_path=source, run_docling=False, run_presidio=False)
    assert result["selected_extractor"] == "deterministic_baseline"
    assert "deterministic_baseline_selected" in str(result["selection_reason"])
    assert result["source"]["original_modified"] is False
    assert source.read_bytes() == before


def test_document_intelligence_and_retrieval_statuses_do_not_report_runtime_downloads() -> None:
    doc_status = document_intelligence_status()
    retrieval_status = optional_backend_status()
    assert doc_status["local_only"] is True
    assert doc_status["network_used"] is False
    assert doc_status["automatic_install"] is False
    assert retrieval_status["automatic_installation"] is False
    assert retrieval_status["automatic_model_download"] is False
    docling = next(row for row in doc_status["adapters"] if row["adapter_id"] == "docling")
    assert docling["mode"] == "isolated_subprocess_offline"
    assert "model" in str(docling["detail"]).lower()
    if docling["available"]:
        assert "expected at" in str(docling["detail"]).lower()


def test_frozen_document_worker_uses_internal_store_command(monkeypatch, tmp_path: Path) -> None:
    from legal.document_intelligence import service

    source = tmp_path / "document.txt"
    source.write_text("private example", encoding="utf-8")
    observed: dict[str, object] = {}

    def fake_run(command, **kwargs):
        observed["command"] = command
        output_path = Path(command[command.index("--document-intelligence-output") + 1])
        output_path.write_text('{"status":"pass","adapter":"presidio","findings":[]}', encoding="utf-8")
        return SimpleNamespace(
            stdout="",
            stderr="",
            returncode=0,
        )

    monkeypatch.setattr(service.sys, "frozen", True, raising=False)
    monkeypatch.setattr(service.sys, "executable", "MaineFamilyLawLLM.exe")
    monkeypatch.setattr(service.subprocess, "run", fake_run)

    result = service._run_worker("presidio", source, timeout=10)

    assert observed["command"] == [
        "MaineFamilyLawLLM.exe",
        "--document-intelligence-worker",
        "presidio",
        str(source),
        "--document-intelligence-output",
        observed["command"][-1],
    ]
    assert result["status"] == "pass"


def test_frozen_document_engines_resolve_beside_executable(monkeypatch, tmp_path: Path) -> None:
    from legal.document_intelligence import service

    runtime_root = tmp_path / "runtime"
    runtime_root.mkdir()
    monkeypatch.setattr(service.sys, "frozen", True, raising=False)
    monkeypatch.setattr(service.sys, "executable", str(runtime_root / "MaineFamilyLawLLM.exe"))

    assert service._bundle_root() == runtime_root
    assert service._bundled_tesseract_root() == runtime_root / "store" / "tesseract"


def test_presidio_worker_disables_public_suffix_network_refresh() -> None:
    source = (Path(__file__).resolve().parents[1] / "legal" / "document_intelligence" / "worker.py").read_text(
        encoding="utf-8"
    )
    assert "tldextract.TLDExtract" in source
    assert "suffix_list_urls=()" in source
    assert "fallback_to_snapshot=True" in source


def test_installed_runtime_api_launches_and_serves_feature_status() -> None:
    resolution = resolve_installed_runtime_executable()
    if not resolution.executable_path:
        pytest.skip("No installed or bundled runtime executable is available.")

    port = _free_port()
    base_url = f"http://127.0.0.1:{port}"
    process = subprocess.Popen(
        [resolution.executable_path, "--serve-local-api", "--port", str(port)],
        cwd=str(Path(resolution.executable_path).parent),
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        text=True,
    )
    try:
        health = _wait_for_json(f"{base_url}/api/health", timeout_s=180)
        version = _wait_for_json(f"{base_url}/api/version", timeout_s=30)
        doc_status = _wait_for_json(f"{base_url}/api/document-intelligence/status", timeout_s=30)
        root = urllib.request.urlopen(f"{base_url}/", timeout=30).read().decode("utf-8", errors="replace")
        # The current canonical local gateway reports ``ok`` for healthy
        # desktop operation; older runtime contracts used ready/degraded.
        assert health["status"] in {"ok", "ready", "degraded"}
        assert resolution.executable_path.endswith("MaineFamilyLawLLM.exe")
        assert version["workbench_url"] == "/"
        assert doc_status["local_only"] is True
        assert "workbench" in root.lower()
    finally:
        try:
            process.terminate()
            process.wait(timeout=30)
        except Exception:
            process.kill()
