from __future__ import annotations

import io
import zipfile
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient

from legal.documents.docx_engine import engine_status, list_docx_paragraphs
from legal.documents.workspace import (
    create_document,
    save_imported_source,
    structured_diff,
    verify_audit_chain,
    workspace_paths,
)
from maine_family_law_llm import api as api_module


def _fictional_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Fictional original paragraph for review.")
    document.add_paragraph("Fictional second paragraph for comment review.")
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_pass71_import_track_comment_reimport_and_compare_through_canonical_api(monkeypatch, tmp_path: Path) -> None:
    root = tmp_path / "fictional-matter"
    root.mkdir()
    monkeypatch.setattr(api_module, "active_case_root", lambda: root)
    client = TestClient(api_module.app)
    assert engine_status()["tracked_changes_available"] is True
    created = client.post(
        "/api/document-workspace/documents",
        json={"title": "Fictional DOCX review", "content": "Fictional original paragraph for review.\nFictional second paragraph for comment review.", "document_type": "draft"},
    )
    assert created.status_code == 200
    document_id = created.json()["document"]["document_id"]
    source_bytes = _fictional_docx_bytes()
    saved = save_imported_source(root, document_id=document_id, data=source_bytes, suffix=".docx")
    assert saved["original_preserved"] is True
    paragraphs = client.get(f"/api/document-workspace/documents/{document_id}/docx/paragraphs?start=1&limit=10")
    assert paragraphs.status_code == 200
    first = paragraphs.json()["paragraphs"][0]
    assert "Fictional original paragraph" in first["text"]
    refused = client.post(
        f"/api/document-workspace/documents/{document_id}/docx/tracked-edit",
        json={"operations": [], "confirmed": False},
    )
    assert refused.status_code == 409
    edited = client.post(
        f"/api/document-workspace/documents/{document_id}/docx/tracked-edit",
        json={
            "confirmed": True,
            "author": "Fictional reviewer",
            "operations": [
                {
                    "action": "replace",
                    "paragraph": first["ref"],
                    "find": "original",
                    "replace_with": "revised",
                    "occurrence": 0,
                },
                {
                    "action": "add_comment",
                    "find": "second paragraph",
                    "comment": "Fictional tracked review comment.",
                    "occurrence": 0,
                },
            ],
        },
    )
    assert edited.status_code == 200
    payload = edited.json()
    assert payload["tracked_changes"] is True and payload["original_preserved"] is True
    assert payload["filing_ready"] is False and payload["review_required"] is True
    artifact = client.get(payload["download_url"])
    assert artifact.status_code == 200 and artifact.content.startswith(b"PK")
    replayed_artifact = client.get(
        payload["download_url"],
        headers={"X-MFLL-Client-Session": "f" * 48},
    )
    assert replayed_artifact.status_code == 404
    assert source_bytes == next(workspace_paths(root).sources.glob(f"{document_id}/original-*.docx")).read_bytes()
    with zipfile.ZipFile(io.BytesIO(artifact.content)) as archive:
        names = set(archive.namelist())
        document_xml = archive.read("word/document.xml")
    assert "word/comments.xml" in names
    assert b"revised" in document_xml.lower()
    reimported_docx = Document(io.BytesIO(artifact.content))
    reimported_text = "\n".join(paragraph.text for paragraph in reimported_docx.paragraphs)
    reimported = create_document(root, title="Fictional reimported tracked copy", content=reimported_text, document_type="draft")
    reimported_source = save_imported_source(root, document_id=reimported["document_id"], data=artifact.content, suffix=".docx")
    assert reimported_source["original_preserved"] is True
    reimported_paragraphs = list_docx_paragraphs(
        source_path=next(workspace_paths(root).sources.glob(f"{reimported['document_id']}/original-*.docx")),
        allowed_source_root=workspace_paths(root).sources,
    )
    assert reimported_paragraphs["paragraphs"]
    comparison = structured_diff(
        "Fictional original paragraph for review.\nFictional second paragraph for comment review.",
        reimported_text,
    )
    assert comparison["changes_count"] >= 1
    assert verify_audit_chain(root)["valid"] is True


def test_pass71_store_build_declares_the_docx_runtime_and_ui_mirrors_the_workflow() -> None:
    root = Path(".")
    requirements = (root / "store/pyinstaller/requirements-store-build.txt").read_text(encoding="utf-8")
    spec = (root / "store/pyinstaller/maine_family_law_llm.spec").read_text(encoding="utf-8")
    src = root / "src/maine_family_law_llm/ui/workbench.js"
    mirror = root / "maine_family_law_llm/ui/workbench.js"
    assert "docx-editor>=0.7.1,<0.8" in requirements
    assert "python-docx>=1.2.0,<2" in requirements
    assert '"docx_editor"' in spec and '"docx-editor"' in spec
    assert src.read_bytes() == mirror.read_bytes()
    text = src.read_text(encoding="utf-8")
    assert "/docx/paragraphs?start=1&limit=300" in text
    assert "/docx/tracked-edit" in text
    assert "Download tracked Word copy" in text
