from __future__ import annotations

import hashlib
import io
import json
import sys
import types
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfWriter

from legal.document_intelligence import service
from legal.document_intelligence.service import (
    DocumentIntelligenceError,
    analyze_document,
    create_ocr_preservation_copy,
    document_intelligence_status,
)
from maine_family_law_llm import api
from maine_family_law_llm.local_workbench_ui import read_workbench_asset, render_local_workbench_html


def _pdf_bytes() -> bytes:
    writer = PdfWriter()
    writer.add_blank_page(width=144, height=144)
    stream = io.BytesIO()
    writer.write(stream)
    return stream.getvalue()


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Motion for Contempt", level=1)
    document.add_paragraph("1. The current order requires weekly contact.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Date"
    table.cell(0, 1).text = "Event"
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _row(case_root: Path, *, evidence_id: str, data: bytes, suffix: str) -> dict:
    staged = case_root / "02_PRIVATE_FORENSIC_MASTER" / "files" / f"{evidence_id}{suffix}"
    staged.parent.mkdir(parents=True, exist_ok=True)
    staged.write_bytes(data)
    return {
        "evidence_id": evidence_id,
        "private_copy_relpath": staged.relative_to(case_root).as_posix(),
        "source_hash": hashlib.sha256(data).hexdigest(),
        "source_type": suffix.lstrip("."),
        "source_locator": staged.name,
        "parser_status": "parsed",
        "text_status": "available",
    }


def _client(monkeypatch: pytest.MonkeyPatch, case_root: Path, rows: list[dict]) -> TestClient:
    monkeypatch.setattr(api, "active_case_root", lambda: case_root)
    monkeypatch.setattr(api, "load_case_search_records", lambda _root: rows)
    api._record_open_tokens.clear()
    api._document_intelligence_artifacts.clear()
    return TestClient(api.app)


def test_v590_status_is_local_optional_and_fail_closed() -> None:
    payload = document_intelligence_status()
    assert payload["local_only"] is True
    assert payload["network_used"] is False
    assert payload["automatic_install"] is False
    by_id = {row["adapter_id"]: row for row in payload["adapters"]}
    assert by_id["deterministic_baseline"]["available"] is True
    assert by_id["docling"]["mode"] == "isolated_subprocess_offline"
    assert by_id["presidio"]["review_required"] is True
    assert by_id["ocrmypdf"]["license"] == "MPL-2.0"


def test_v590_baseline_docx_structure_and_privacy_receipt(tmp_path: Path) -> None:
    case_root = tmp_path / "case"
    case_root.mkdir()
    source = case_root / "motion.docx"
    source.write_bytes(_docx_bytes())
    before = source.read_bytes()
    result = analyze_document(
        case_root=case_root,
        source_path=source,
        run_docling=False,
        run_presidio=False,
    )
    assert result["selected_extractor"] == "deterministic_baseline"
    assert result["structured_document"]["block_count"] >= 3
    kinds = {row["kind"] for row in result["structured_document"]["blocks"]}
    assert "table_row" in kinds
    assert result["source"]["original_modified"] is False
    assert source.read_bytes() == before
    artifact = case_root / result["artifact"]["relative_path"]
    assert artifact.is_file()
    assert hashlib.sha256(artifact.read_bytes()).hexdigest() == result["artifact"]["sha256"]
    stored = json.loads(artifact.read_text(encoding="utf-8"))
    assert stored["receipt_sha256"] == result["receipt_sha256"]


def test_v590_privacy_review_returns_spans_not_plain_sensitive_values(tmp_path: Path) -> None:
    case_root = tmp_path / "case"
    case_root.mkdir()
    source = case_root / "record.txt"
    source.write_text(
        "DOB: 01/02/2010\nEmail parent@example.com\nCall (207) 555-1212\nSSN 123-45-6789",
        encoding="utf-8",
    )
    result = analyze_document(case_root=case_root, source_path=source, run_docling=False, run_presidio=False)
    privacy = result["privacy_review"]
    assert privacy["finding_counts"]["DATE_OF_BIRTH"] == 1
    assert privacy["finding_counts"]["EMAIL_ADDRESS"] == 1
    assert privacy["finding_counts"]["PHONE_NUMBER"] == 1
    assert privacy["finding_counts"]["US_SSN"] == 1
    serialized = json.dumps(privacy)
    assert "parent@example.com" not in serialized
    assert "123-45-6789" not in serialized
    assert privacy["complete_detection_guaranteed"] is False


def test_v590_docling_and_presidio_results_are_optional_layers(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    case_root = tmp_path / "case"
    case_root.mkdir()
    source = case_root / "record.txt"
    source.write_text("Heading\nA private statement.", encoding="utf-8")

    real_status = service.document_intelligence_status()
    for row in real_status["adapters"]:
        if row["adapter_id"] in {"docling", "presidio"}:
            row["available"] = True
    monkeypatch.setattr(service, "document_intelligence_status", lambda: real_status)

    def fake_worker(adapter: str, path: Path, *, timeout: int) -> dict:
        if adapter == "docling":
            return {
                "status": "pass",
                "adapter": "docling",
                "blocks": [{"block_id": "docling_1", "kind": "heading", "text": "Heading", "page_number": 1, "order": 1}],
            }
        return {
            "status": "pass",
            "adapter": "presidio",
            "findings": [{"entity_type": "PERSON", "start": 0, "end": 7, "score": 0.8, "recognizer": "presidio"}],
        }

    monkeypatch.setattr(service, "_run_worker", fake_worker)
    result = analyze_document(case_root=case_root, source_path=source, run_docling=True, run_presidio=True)
    assert result["selected_extractor"] == "docling"
    assert result["structured_document"]["blocks"][0]["block_id"].startswith("blk_")
    assert result["privacy_review"]["presidio_status"] == "pass"
    assert "presidio" in result["privacy_review"]["detectors"]


def test_v590_refuses_outside_symlink_oversize_and_hash_change(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    case_root = tmp_path / "case"
    case_root.mkdir()
    outside = tmp_path / "outside.txt"
    outside.write_text("outside")
    with pytest.raises(DocumentIntelligenceError) as exc:
        analyze_document(case_root=case_root, source_path=outside, run_docling=False, run_presidio=False)
    assert exc.value.code == "document_source_outside_matter"

    source = case_root / "inside.txt"
    source.write_text("inside")
    with pytest.raises(DocumentIntelligenceError) as exc:
        analyze_document(case_root=case_root, source_path=source, source_hash="0" * 64, run_docling=False, run_presidio=False)
    assert exc.value.code == "document_source_hash_mismatch"

    link = case_root / "linked.txt"
    try:
        link.symlink_to(source)
    except OSError:
        pytest.skip("symlinks unavailable")
    with pytest.raises(DocumentIntelligenceError) as exc:
        analyze_document(case_root=case_root, source_path=link, run_docling=False, run_presidio=False)
    assert exc.value.code == "document_source_invalid"

    monkeypatch.setattr(service, "MAX_SOURCE_BYTES", 2)
    with pytest.raises(DocumentIntelligenceError) as exc:
        analyze_document(case_root=case_root, source_path=source, run_docling=False, run_presidio=False)
    assert exc.value.code == "document_source_too_large"


def test_v590_ocr_preservation_requires_consent_and_never_overwrites(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    case_root = tmp_path / "case"
    case_root.mkdir()
    source = case_root / "scan.pdf"
    source.write_bytes(_pdf_bytes())
    before = source.read_bytes()
    with pytest.raises(DocumentIntelligenceError) as exc:
        create_ocr_preservation_copy(case_root=case_root, source_path=source, approved=False)
    assert exc.value.code == "ocr_preservation_consent_required"

    monkeypatch.setattr(service, "local_ocr_engine_status", lambda: {"available": False, "pdf_ocr_available": False})
    blocked = create_ocr_preservation_copy(case_root=case_root, source_path=source, approved=True)
    assert blocked["status"] == "blocked"
    assert blocked["blockers"] == ["ocrmypdf_not_installed"]
    assert source.read_bytes() == before


def test_v590_ocr_preservation_success_creates_separate_hash_bound_copy(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    case_root = tmp_path / "case"
    case_root.mkdir()
    source = case_root / "scan.pdf"
    source.write_bytes(_pdf_bytes())
    before = source.read_bytes()
    original_module_available = service._module_available
    monkeypatch.setattr(service, "local_ocr_engine_status", lambda: {"available": True, "pdf_ocr_available": True})
    monkeypatch.setattr(
        service,
        "_module_available",
        lambda module: True if module == "ocrmypdf" else original_module_available(module),
    )
    fake_api = types.ModuleType("ocrmypdf.api")

    def fake_ocr(input_file, output_file, **kwargs):
        sidecar = Path(kwargs["sidecar"])
        output_path = Path(output_file)
        output_path.write_bytes(Path(input_file).read_bytes() + b"\nDERIVED")
        sidecar.write_text("recognized text", encoding="utf-8")
        return 0

    fake_api.ocr = fake_ocr  # type: ignore[attr-defined]
    fake_package = types.ModuleType("ocrmypdf")
    fake_package.__path__ = []  # type: ignore[attr-defined]
    fake_package.api = fake_api  # type: ignore[attr-defined]
    monkeypatch.setitem(sys.modules, "ocrmypdf", fake_package)
    monkeypatch.setitem(sys.modules, "ocrmypdf.api", fake_api)
    result = create_ocr_preservation_copy(case_root=case_root, source_path=source, approved=True)
    assert result["status"] == "pass"
    assert result["original_modified"] is False
    assert source.read_bytes() == before
    output = case_root / result["artifacts"]["pdf"]["relative_path"]
    assert output.is_file()
    assert output != source
    assert hashlib.sha256(output.read_bytes()).hexdigest() == result["output_sha256"]


def test_v590_api_uses_opaque_record_and_artifact_tokens(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    case_root = tmp_path / "case"
    row = _row(case_root, evidence_id="REC-1", data=b"MOTION FOR CONTEMPT\nEmail: parent@example.com", suffix=".txt")
    client = _client(monkeypatch, case_root, [row])
    token = api._record_open_token(case_root, "REC-1", "REC-1.txt")

    status = client.get("/api/document-intelligence/status")
    assert status.status_code == 200
    assert status.json()["local_only"] is True

    denied = client.post("/api/document-intelligence/analyze", json={"source_token": token, "approved": False})
    assert denied.status_code == 409

    analyzed = client.post(
        "/api/document-intelligence/analyze",
        json={"source_token": token, "approved": True, "run_docling": False, "run_presidio": False},
    )
    assert analyzed.status_code == 200
    payload = analyzed.json()
    assert payload["source"]["filename"] == "REC-1.txt"
    assert "relative_path" not in payload["artifact"]
    assert len(payload["artifact"]["download_token"]) == 64
    serialized = json.dumps(payload)
    assert str(case_root) not in serialized

    download = client.get(payload["artifact"]["download_url"])
    assert download.status_code == 200
    assert download.headers["x-mfl-hash-verified"] == "true"

    other = tmp_path / "other"
    other.mkdir()
    monkeypatch.setattr(api, "active_case_root", lambda: other)
    assert client.get(payload["artifact"]["download_url"]).status_code == 404


def test_v590_ui_exposes_document_intelligence_in_verified_source_inspector() -> None:
    html = render_local_workbench_html()
    js = read_workbench_asset("workbench.js")
    css = read_workbench_asset("workbench.css")
    for marker in (
        'id="record-inspector-document-intelligence"',
        'id="document-intelligence-modal"',
        'id="document-intelligence-analyze"',
        'id="document-intelligence-ocr"',
    ):
        assert marker in html
    for marker in (
        "/api/document-intelligence/status",
        "/api/records/",
        "function runDocumentIntelligencePrivacyScan",
        "function runDocumentIntelligenceRedactedCopy",
        "function renderDocumentIntelligenceReport",
        "function runDocumentIntelligenceOcr",
    ):
        assert marker in js
    assert ".document-intelligence-modal" in css
    assert "original remains immutable" in html


def test_v590_record_scoped_document_intelligence_routes_cover_integrity_blocks_and_privacy(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    case_root = tmp_path / "case"
    case_root.mkdir()
    source = case_root / "record.docx"
    source.write_bytes(_docx_bytes())
    row = _row(case_root, evidence_id="REC-1", data=source.read_bytes(), suffix=".docx")
    client = _client(monkeypatch, case_root, [row])

    integrity = client.get("/api/records/REC-1/integrity")
    assert integrity.status_code == 200
    integrity_payload = integrity.json()
    assert integrity_payload["record_id"] == "REC-1"
    assert integrity_payload["integrity"]["immutable_original"] is True
    assert str(case_root) not in json.dumps(integrity_payload)

    blocks = client.get("/api/records/REC-1/blocks?limit=5&offset=0")
    assert blocks.status_code == 200
    block_payload = blocks.json()
    assert block_payload["record_id"] == "REC-1"
    assert block_payload["total"] >= 1
    assert block_payload["blocks"]

    parsed = client.post(
        "/api/records/REC-1/parse",
        json={"source_token": api._record_open_token(case_root, "REC-1", "record.docx"), "approved": True, "run_docling": False, "run_presidio": False},
    )
    assert parsed.status_code == 200
    parsed_payload = parsed.json()
    assert parsed_payload["record_id"] == "REC-1"
    assert parsed_payload["structured_document"]["block_count"] >= 1
    assert parsed_payload["artifact"]["download_url"].startswith("/api/document-intelligence/artifacts/")

    privacy = client.post(
        "/api/records/REC-1/privacy-scan",
        json={"source_token": api._record_open_token(case_root, "REC-1", "record.docx"), "approved": True, "run_presidio": False},
    )
    assert privacy.status_code == 200
    privacy_payload = privacy.json()
    assert privacy_payload["record_id"] == "REC-1"
    assert privacy_payload["privacy_review"]["review_required"] is True


def test_v590_record_duplicates_compare_and_redaction_receipts(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    case_root = tmp_path / "case"
    case_root.mkdir()
    doc_bytes = _docx_bytes()
    row_a = _row(case_root, evidence_id="REC-1", data=doc_bytes, suffix=".docx")
    row_b = _row(case_root, evidence_id="REC-2", data=doc_bytes, suffix=".docx")
    row_b["text_excerpt"] = "Some other text with parent@example.com"
    rows = [row_a, row_b]
    client = _client(monkeypatch, case_root, rows)

    duplicates = client.get("/api/records/REC-1/duplicates")
    assert duplicates.status_code == 200
    dup_payload = duplicates.json()
    assert dup_payload["exact_duplicate"] is True
    assert len(dup_payload["exact_duplicates"]) == 2

    compare = client.post("/api/records/compare", json={"left_record_id": "REC-1", "right_record_id": "REC-2"})
    assert compare.status_code == 200
    compare_payload = compare.json()
    assert compare_payload["left"]["record_id"] == "REC-1"
    assert compare_payload["right"]["record_id"] == "REC-2"

    redaction = client.post(
        "/api/records/REC-2/redacted-copy",
        json={"source_token": api._record_open_token(case_root, "REC-2", "record.docx"), "approved": True, "reviewer": "tester", "run_presidio": False},
    )
    assert redaction.status_code == 200
    redaction_payload = redaction.json()
    redacted = redaction_payload["artifacts"]["redacted_copy"]
    receipt = client.get(redaction_payload["artifacts"]["receipt"]["download_url"])
    assert receipt.status_code == 200
    assert receipt.json()["artifact_type"] == "redacted_copy"
    redacted_receipt = client.get(f"/api/artifacts/{redacted['artifact_id']}/receipt")
    assert redacted_receipt.status_code == 200
    assert redacted_receipt.json()["artifact_type"] == "redacted_copy"


def test_v590_ocr_derivative_receipt_is_separate_and_hash_bound(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    case_root = tmp_path / "case"
    case_root.mkdir()
    source = case_root / "scan.pdf"
    source.write_bytes(_pdf_bytes())
    row = _row(case_root, evidence_id="REC-1", data=source.read_bytes(), suffix=".pdf")
    client = _client(monkeypatch, case_root, [row])

    original_module_available = service._module_available
    monkeypatch.setattr(service, "local_ocr_engine_status", lambda: {"available": True, "pdf_ocr_available": True})
    monkeypatch.setattr(
        service,
        "_module_available",
        lambda module: True if module == "ocrmypdf" else original_module_available(module),
    )
    fake_api = types.ModuleType("ocrmypdf.api")

    def fake_ocr(input_file, output_file, **kwargs):
        sidecar = Path(kwargs["sidecar"])
        output_path = Path(output_file)
        output_path.write_bytes(Path(input_file).read_bytes() + b"\nDERIVED")
        sidecar.write_text("recognized text", encoding="utf-8")
        return 0

    fake_api.ocr = fake_ocr  # type: ignore[attr-defined]
    fake_package = types.ModuleType("ocrmypdf")
    fake_package.__path__ = []  # type: ignore[attr-defined]
    fake_package.api = fake_api  # type: ignore[attr-defined]
    monkeypatch.setitem(sys.modules, "ocrmypdf", fake_package)
    monkeypatch.setitem(sys.modules, "ocrmypdf.api", fake_api)

    ocr = client.post(
        "/api/records/REC-1/ocr",
        json={"source_token": api._record_open_token(case_root, "REC-1", "scan.pdf"), "approved": True, "language": "eng"},
    )
    assert ocr.status_code == 200
    ocr_payload = ocr.json()
    pdf_artifact = ocr_payload["artifacts"]["pdf"]
    receipt = client.get(pdf_artifact["receipt_url"])
    assert receipt.status_code == 200
    assert receipt.json()["artifact_type"] == "ocr_preservation_pdf"
